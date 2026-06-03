#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys, json, os
from anarci import anarci
from Bio import SeqIO
from docx import Document
from docx.shared import RGBColor

# CDR 定义（IMGT）
CDR_DEFS = {
    "CDR1": range(27, 39),
    "CDR2": range(56, 66),
    "CDR3": range(105, 118),
}

# 终端颜色
ANSI = {
    "CDR1": "\033[31m",
    "CDR2": "\033[32m",
    "CDR3": "\033[34m",
    "RESET": "\033[0m"
}

# Word 颜色
CDR_COLORS = {
    "CDR1": RGBColor(255, 0, 0),
    "CDR2": RGBColor(0, 128, 0),
    "CDR3": RGBColor(0, 0, 255)
}

# ----------------------------
def map_cdrs_to_original(numbering_list, original_seq):
    # numbering_list: [( (number, ins), aa ), ...]
    map_align_to_orig = []
    orig_counter = 0
    for pos, aa in numbering_list:
        if aa == '-' or aa is None:
            map_align_to_orig.append(None)
        else:
            orig_counter += 1
            map_align_to_orig.append(orig_counter)

    cdrs = {}
    for name, imgt_range in CDR_DEFS.items():
        aa_list = []
        start_orig = None
        end_orig = None
        for i, (pos, aa) in enumerate(numbering_list):
            if pos is None or aa in (None,'-'):
                continue
            num, ins = pos
            if num in imgt_range:
                mapped = map_align_to_orig[i]
                if mapped is not None:
                    if start_orig is None:
                        start_orig = mapped
                    end_orig = mapped
                    aa_list.append(aa)
        cdrs[name] = {"seq": "".join(aa_list), "start": start_orig, "end": end_orig} if aa_list else {"seq": "", "start": None, "end": None}
    return cdrs

def colorize_sequence(original_seq, cdr_info):
    n = len(original_seq)
    tag_map = [None]*n
    for name in ("CDR1","CDR2","CDR3"):
        info = cdr_info.get(name,{})
        s = info.get("start")
        e = info.get("end")
        if s and e:
            for i in range(s-1, e):
                if 0 <= i < n:
                    tag_map[i] = name
    # 终端彩色序列
    colored_chars = []
    current_tag = None
    for i, aa in enumerate(original_seq):
        tag = tag_map[i]
        if tag != current_tag:
            if current_tag is not None:
                colored_chars.append(ANSI["RESET"])
            if tag is not None:
                colored_chars.append(ANSI[tag])
            current_tag = tag
        colored_chars.append(aa)
    if current_tag is not None:
        colored_chars.append(ANSI["RESET"])
    colored_seq = "".join(colored_chars)

    # 带标签的序列
    tagged_chars = []
    current_tag = None
    for i, aa in enumerate(original_seq):
        tag = tag_map[i]
        if tag != current_tag:
            if current_tag is not None:
                tagged_chars.append(f"[/{current_tag}]")
            if tag is not None:
                tagged_chars.append(f"[{tag}:{i+1}]")
            current_tag = tag
        tagged_chars.append(aa)
    if current_tag is not None:
        tagged_chars.append(f"[/{current_tag}]")
    tagged_seq = "".join(tagged_chars)
    return colored_seq, tagged_seq

def write_word(fasta_file, cdr_data, output_docx):
    doc = Document()
    for rec in SeqIO.parse(fasta_file, "fasta"):
        seq_id = rec.id
        seq = str(rec.seq)
        doc.add_paragraph(f">{seq_id}", style='Intense Quote')
        p = doc.add_paragraph()
        cdr_info = cdr_data.get(seq_id, {}).get("CDRs",{})
        for i, aa in enumerate(seq):
            run = p.add_run(aa)
            color_set = False
            for cdr_name in ("CDR1","CDR2","CDR3"):
                info = cdr_info.get(cdr_name,{})
                s = info.get("start")
                e = info.get("end")
                if s and e and (s-1) <= i <= (e-1):
                    run.font.color.rgb = CDR_COLORS[cdr_name]
                    color_set = True
                    break
            if not color_set:
                run.font.color.rgb = RGBColor(0,0,0)
    doc.save(output_docx)
    print(f"Word 文件生成: {output_docx}", file=sys.stderr)

def process_fasta(fasta_file):
    results = {}
    base = os.path.splitext(os.path.basename(fasta_file))[0]
    colored_fname = f"{base}_colored.fasta"
    tagged_fname = f"{base}_tagged.fasta"
    word_fname = f"{base}_colored.docx"

    with open(colored_fname,"w",encoding="utf-8") as cf, open(tagged_fname,"w",encoding="utf-8") as tf:
        for rec in SeqIO.parse(fasta_file,"fasta"):
            seq = str(rec.seq).strip()
            if not seq:
                continue
            try:
                res = anarci([("seq", seq)], scheme="imgt")
            except Exception as e:
                print(f"ANARCI 识别失败: {e}", file=sys.stderr)
                results[rec.id] = {}
                continue
            try:
                numbering_list = res[0][0][0][0]  # [( (num, ins), aa ), ...]
                chain_type = res[1][0][0].get("chain_type","") if res[1] and res[1][0] else ""
            except Exception as e:
                print(f"解析 ANARCI 返回值失败: {e}", file=sys.stderr)
                results[rec.id] = {}
                continue

            cdr_info = map_cdrs_to_original(numbering_list, seq)
            results[rec.id] = {"chain_type": chain_type, "sequence_length": len(seq), "CDRs": cdr_info}

            colored_seq, tagged_seq = colorize_sequence(seq, cdr_info)
            # 彩色FASTA
            cf.write(f">{rec.id} colored\n")
            for i in range(0,len(colored_seq),60):
                cf.write(colored_seq[i:i+60]+"\n")
            # 标签FASTA
            tf.write(f">{rec.id} tagged\n")
            for i in range(0,len(tagged_seq),60):
                tf.write(tagged_seq[i:i+60]+"\n")

    write_word(fasta_file, results, word_fname)
    return results, colored_fname, tagged_fname, word_fname

def main():
    if len(sys.argv)<2:
        print("用法: python3 fab_cdr_export.py input.fasta > output.json")
        sys.exit(1)
    fasta_file = sys.argv[1]
    results, colored_file, tagged_file, word_file = process_fasta(fasta_file)
    print(json.dumps(results, ensure_ascii=False, indent=2))
    print(f"彩色FASTA: {colored_file}", file=sys.stderr)
    print(f"带标签FASTA: {tagged_file}", file=sys.stderr)
    print(f"Word文件: {word_file}", file=sys.stderr)

if __name__=="__main__":
    main()
